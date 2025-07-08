from flask import Flask, render_template, request, redirect, session, flash, send_file
import os
from werkzeug.utils import secure_filename
from resume_parser import extract_text
from score_calculator import calculate_match_score
from database import add_user, get_user_by_email, init_db
from docx import Document
from dotenv import load_dotenv
from flask import Flask, render_template, request, jsonify




# Load environment variables
load_dotenv()
app = Flask(__name__)
app.secret_key = 'surya_resume_ai_2025_secure_key'

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ========== ROUTES ========== #

@app.route('/')
def index():
    return render_template('index.html')


# ---------- Signup ----------
@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        name = request.form['name']
        email = request.form['email']
        password = request.form['password']
        existing_user = get_user_by_email(email)
        if existing_user:
            flash('User already exists!')
            return redirect('/signup')
        add_user(name, email, password)
        flash('Account created. Please login.')
        return redirect('/login')
    return render_template('signup.html')


# ---------- Login ----------
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        user = get_user_by_email(email)
        if user and user[3] == password:
            session['user'] = user[1]
            flash('Login successful!')
            return redirect('/dashboard')
        flash('Invalid credentials.')
    return render_template('login.html')


# ---------- Logout ----------
@app.route('/logout')
def logout():
    session.pop('user', None)
    flash('Logged out successfully.')
    return redirect('/')


# ---------- Dashboard ----------
@app.route('/dashboard', methods=['GET'])
def dashboard():
    if 'user' not in session:
        return redirect('/login')
    return render_template('dashboard.html')


# ---------- Analyze Resume ----------
@app.route('/analyze', methods=['POST'])
def analyze():
    if 'user' not in session:
        return redirect('/login')

    # Get uploaded files
    resume_file = request.files['resume']
    jd_file = request.files['jd']

    # Validate
    if not resume_file or not jd_file:
        flash('Please upload both resume and job description files.')
        return redirect('/dashboard')

    # Save files
    resume_path = os.path.join(UPLOAD_FOLDER, secure_filename(resume_file.filename))
    jd_path = os.path.join(UPLOAD_FOLDER, secure_filename(jd_file.filename))
    resume_file.save(resume_path)
    jd_file.save(jd_path)

    # Extract text using resume_parser.py
    resume_text = extract_text(resume_path)
    jd_text = extract_text(jd_path)

    # Calculate score using score_calculator.py
    score = calculate_match_score(resume_text, jd_text)

    return render_template('result.html', score=score)


# ---------- Resume Builder ----------
@app.route('/builder', methods=['GET'])
def builder():
    if 'user' not in session:
        return redirect('/login')
    return render_template('builder.html')


@app.route('/build_resume', methods=['POST'])
def build_resume():
    name = request.form.get('name')
    email = request.form.get('email')
    summary = request.form.get('summary')
    skills = request.form.get('skills')
    experience = request.form.get('experience')

    doc = Document()
    doc.add_heading(name, 0)
    doc.add_paragraph(email)
    doc.add_heading('Profesqsional Summary', level=1)
    doc.add_paragraph(summary)
    doc.add_heading('Skills', level=1)
    doc.add_paragraph(skills)
    doc.add_heading('Experience', level=1)
    doc.add_paragraph(experience)

    output_path = os.path.join(UPLOAD_FOLDER, 'generated_resume.docx')
    doc.save(output_path)

    return send_file(output_path, as_attachment=True)


# ---------- Template Gallery ----------
@app.route('/templates', methods=['GET'])
def templates_gallery():
    return render_template('templates_gallery.html')


# ---------- Download Summary ----------
@app.route('/download_summary', methods=['GET'])
def download_summary():
    if 'user' not in session:
        return redirect('/login')
    
    score = request.args.get('score', '0')
    summary_text = f"📄 Resume Match Summary\n\nMatch Score: {score}%\n\nThank you for using ResumeXpert!"

    summary_path = os.path.join('uploads', 'summary.txt')
    with open(summary_path, 'w', encoding='utf-8') as f:
        f.write(summary_text)

    return send_file(summary_path, as_attachment=True)


# ---------- Download Template ----------
@app.route('/download_template/<template_name>', methods=['GET'])
def download_template(template_name):
    try:
        file_path = os.path.join('templates', f'{template_name}.docx')
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
        else:
            return f"❌ File not found: {template_name}.docx"
    except Exception as e:
        return f"⚠️ Error: {str(e)}"


# ---------- Run the App ----------
if __name__ == '__main__':
    init_db()
    app.run(host='0.0.0.0', port=5000, debug=True)

